"""Generate DOCX file from an Article — with proper page header, figures,
tables, and a locked 'OpenJATS Body' paragraph style for clean justify."""
import base64
from io import BytesIO
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from models import Article, Figure, Journal
from citation_formatter import format_references


# ---------- Setup helpers ----------

def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        # Header / footer distance
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)


def _suppress_hyphens(paragraph):
    """Disable auto-hyphenation on a paragraph."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    sup = OxmlElement("w:suppressAutoHyphens")
    sup.set(qn("w:val"), "true")
    pPr.append(sup)


def _register_styles(doc: Document):
    """Register custom paragraph styles for body text, headings, captions."""
    styles = doc.styles

    # Body style — justified, line 1.15, no hyphenation, indent 0.5cm
    if "OpenJATS Body" not in styles:
        body = styles.add_style("OpenJATS Body", WD_STYLE_TYPE.PARAGRAPH)
        body.base_style = styles["Normal"]
        body.font.name = "Times New Roman"
        body.font.size = Pt(11)
        pf = body.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.5)
        pf.line_spacing = 1.15
        pf.space_after = Pt(4)
        # Disable hyphenation at style level
        pPr = body.element.get_or_add_pPr()
        sup = OxmlElement("w:suppressAutoHyphens")
        sup.set(qn("w:val"), "true")
        pPr.append(sup)

    if "OpenJATS Abstract" not in styles:
        absty = styles.add_style("OpenJATS Abstract", WD_STYLE_TYPE.PARAGRAPH)
        absty.base_style = styles["Normal"]
        absty.font.name = "Times New Roman"
        absty.font.size = Pt(10)
        pf = absty.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.15
        pf.space_after = Pt(4)
        pPr = absty.element.get_or_add_pPr()
        sup = OxmlElement("w:suppressAutoHyphens")
        sup.set(qn("w:val"), "true")
        pPr.append(sup)

    if "OpenJATS H1" not in styles:
        h1 = styles.add_style("OpenJATS H1", WD_STYLE_TYPE.PARAGRAPH)
        h1.font.name = "Calibri"
        h1.font.size = Pt(13)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        pf = h1.paragraph_format
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.keep_with_next = True

    if "OpenJATS H2" not in styles:
        h2 = styles.add_style("OpenJATS H2", WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = "Calibri"
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x33, 0x40, 0x55)
        pf = h2.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)
        pf.keep_with_next = True

    if "OpenJATS Caption" not in styles:
        cap = styles.add_style("OpenJATS Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.base_style = styles["Normal"]
        cap.font.name = "Calibri"
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        pf = cap.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(6)
        pf.space_before = Pt(2)

    if "OpenJATS Reference" not in styles:
        ref = styles.add_style("OpenJATS Reference", WD_STYLE_TYPE.PARAGRAPH)
        ref.base_style = styles["Normal"]
        ref.font.name = "Times New Roman"
        ref.font.size = Pt(10)
        pf = ref.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.15
        pf.left_indent = Cm(0.7)
        pf.first_line_indent = Cm(-0.7)
        pf.space_after = Pt(3)
        pPr = ref.element.get_or_add_pPr()
        sup = OxmlElement("w:suppressAutoHyphens")
        sup.set(qn("w:val"), "true")
        pPr.append(sup)


def _set_page_borders(doc: Document):
    """Add a subtle bottom border under the page header."""
    pass  # handled per paragraph if needed


def _decode_data_url(data_url: str) -> bytes:
    if not data_url:
        return b""
    data = data_url.split(",", 1)[1] if "," in data_url else data_url
    try:
        return base64.b64decode(data)
    except Exception:
        return b""


def _build_page_header(doc: Document, j: Journal, article: Article):
    """Construct a locked page header with logo + journal title + meta info.
    Repeats on every page; not part of body content."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    # Clear any default
    for p in list(header.paragraphs):
        p.clear()

    # First paragraph: logo + journal title (two-column-ish via tab)
    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Embed raster logo
    if j.logo and "svg" not in (j.logo or "").lower():
        img_bytes = _decode_data_url(j.logo)
        if img_bytes:
            run = p1.add_run()
            try:
                run.add_picture(BytesIO(img_bytes), height=Cm(1.2))
            except Exception:
                pass
            p1.add_run("  ")  # spacing after logo

    title_run = p1.add_run(j.title or "")
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Second paragraph: custom header + meta
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bits = []
    if j.custom_header:
        bits.append(j.custom_header)
    meta_parts = []
    if j.volume or j.issue or j.year:
        meta_parts.append(f"Vol. {j.volume or '—'}, No. {j.issue or '—'} ({j.year or '—'})")
    if j.issn:
        meta_parts.append(f"ISSN: {j.issn}")
    if j.eissn:
        meta_parts.append(f"e-ISSN: {j.eissn}")
    if meta_parts:
        bits.append("  ·  ".join(meta_parts))
    if bits:
        meta_run = p2.add_run("  ·  ".join(bits))
        meta_run.font.size = Pt(8)
        meta_run.font.name = "Calibri"
        meta_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        meta_run.italic = True

    # Bottom border under header
    pPr = p2.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "475569")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_page_footer(doc: Document):
    """Build a footer with page number  '1 / N'."""
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        return fld

    fp._p.append(_field("PAGE"))
    sep_run = fp.add_run(" / ")
    sep_run.font.size = Pt(9)
    sep_run.font.name = "Calibri"
    fp._p.append(_field("NUMPAGES"))


# ---------- Body helpers ----------

def _heading(doc: Document, text: str, level: int = 1):
    style_name = "OpenJATS H1" if level == 1 else "OpenJATS H2"
    p = doc.add_paragraph(text, style=style_name)
    return p


def _add_body_para(doc: Document, text: str, style: str = "OpenJATS Body", indent: bool = True, font_family: str = "Times New Roman"):
    p = doc.add_paragraph(style=style)
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    # Set base font on style override
    _add_inline_runs(p, text, base_font=font_family, base_size=11)
    return p


def _add_list(doc: Document, items: list, ordered: bool, font_family: str = "Times New Roman"):
    for idx, raw in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        prefix = f"{idx}." if ordered else "•"
        # Marker run
        mr = p.add_run(f"{prefix}  ")
        mr.font.size = Pt(11)
        mr.font.name = font_family
        # Item content with inline formatting
        _add_inline_runs(p, raw, base_font=font_family, base_size=11)


def _add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_paragraph_shading(paragraph, color_hex: str):
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _set_paragraph_left_border(paragraph, color_hex: str = "475569", size: int = 24):
    """Add a thick left border to mimic the PDF abstract bar."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_inline_runs(paragraph, text: str, base_font: str = "Times New Roman", base_size: int = 11):
    """Parse markdown-ish inline: **bold**, *italic*, _underline_, and add runs."""
    # tokens: split on ** first, then * inside non-bold, then _ inside both
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        bold = italic = underline = False
        seg = part
        if part.startswith("**") and part.endswith("**"):
            bold = True
            seg = part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            italic = True
            seg = part[1:-1]
        elif part.startswith("_") and part.endswith("_"):
            underline = True
            seg = part[1:-1]
        run = paragraph.add_run(seg)
        run.font.name = base_font
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = italic
        run.underline = underline


def _shade_cell(cell, color_hex: str = "E2E8F0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm) and widths_cm[i] > 0:
                cell.width = Cm(widths_cm[i])


def _parse_table_rows(lines: list) -> list:
    rows = []
    for ln in lines:
        if not ln.strip().startswith("|"):
            continue
        stripped = ln.strip()
        if re.match(r"^\|\s*-+\s*(\|\s*-+\s*)*\|?$", stripped):
            continue
        cells = stripped.strip("|").split("|")
        rows.append([c.strip() for c in cells])
    return rows


def _add_table(doc: Document, label: str, caption: str, rows: list, widths_cm=None):
    if not rows or not rows[0]:
        return
    if label or caption:
        cap = doc.add_paragraph(style="OpenJATS Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if label:
            r = cap.add_run(f"{label}. ")
            r.bold = True
            r.italic = False
        cap.add_run(caption)
        cap.paragraph_format.keep_with_next = True
        cap.paragraph_format.space_after = Pt(2)

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(cols):
            value = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            if r_idx == 0:
                run.bold = True
                _shade_cell(cell, "E2E8F0")
    if widths_cm:
        _set_col_widths(table, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_figure(doc: Document, fig: Figure):
    img_bytes = _decode_data_url(fig.data_url)
    if not img_bytes or "svg" in (fig.data_url or "").lower():
        p = doc.add_paragraph(style="OpenJATS Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{fig.label}: image not embeddable]")
        run.italic = True
        return
    width = Cm(14) if fig.width == "full" else Cm(8) if fig.width == "half" else Cm(5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        doc.add_picture(BytesIO(img_bytes), width=width)
    except Exception:
        pass
    cap = doc.add_paragraph(style="OpenJATS Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = cap.add_run(f"{fig.label}. ")
    label_run.bold = True
    label_run.italic = False
    cap.add_run(fig.caption)


def _render_body(doc: Document, body: str, figures_map: dict, font_family: str = "Times New Roman"):
    if not body or not body.strip():
        return
    lines = body.split("\n")
    i = 0
    para_buf = []

    def flush():
        if not para_buf:
            return
        text = " ".join(para_buf).strip()
        para_buf.clear()
        if text:
            _add_body_para(doc, text, font_family=font_family)

    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # Numbered or bullet list
        if re.match(r"^(?:\d+\.|-|\*)\s+", stripped) and not re.match(r"^[-*]{3,}$", stripped):
            flush()
            ordered = bool(re.match(r"^\d+\.\s+", stripped))
            items = []
            while i < len(lines) and re.match(r"^(?:\d+\.|-|\*)\s+", lines[i].strip()) and not re.match(r"^[-*]{3,}$", lines[i].strip()):
                t = lines[i].strip()
                content = re.sub(r"^(?:\d+\.|-|\*)\s+", "", t)
                items.append(content)
                i += 1
            _add_list(doc, items, ordered, font_family=font_family)
            continue

        if re.match(r"^\[PAGE\s*BREAK\]$", stripped, re.IGNORECASE):
            flush()
            _add_page_break(doc)
            i += 1
            continue

        m_fig = re.match(r"^\[Figure:([^\]]+)\]\s*$", stripped)
        if m_fig:
            flush()
            fig = figures_map.get(m_fig.group(1).strip())
            if fig:
                _add_figure(doc, fig)
            i += 1
            continue

        m_tbl = re.match(r"^\[Table:([^\]]+)\]\s*(.*)?$", stripped)
        widths_marker = None
        if m_tbl or stripped.startswith("|"):
            flush()
            label = ""
            caption = ""
            if m_tbl:
                label = m_tbl.group(1).strip()
                caption = (m_tbl.group(2) or "").strip()
                i += 1
            # Read optional widths line: [widths: 3, 4, 5] (cm) or 20%,40%,40%
            if i < len(lines) and re.match(r"^\s*\[widths:[^\]]+\]\s*$", lines[i], re.IGNORECASE):
                wmatch = re.match(r"^\s*\[widths:\s*([^\]]+)\s*\]\s*$", lines[i], re.IGNORECASE)
                if wmatch:
                    widths_marker = wmatch.group(1)
                i += 1
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            parsed = _parse_table_rows(tbl_lines)
            widths_cm = None
            if widths_marker and parsed:
                cols = len(parsed[0])
                widths_cm = _resolve_widths(widths_marker, cols, page_width_cm=16.0)
            if parsed:
                _add_table(doc, label, caption, parsed, widths_cm=widths_cm)
            continue

        if stripped.startswith("### "):
            flush()
            _heading(doc, stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            flush()
            _heading(doc, stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) >= 4:
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:-2].strip())
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            i += 1
            continue

        if not stripped:
            flush()
            i += 1
            continue
        para_buf.append(stripped)
        i += 1
    flush()


def _resolve_widths(spec: str, cols: int, page_width_cm: float = 16.0):
    """Parse '20%, 40%, 40%' or '3, 4, 5' into list of cm widths."""
    parts = [p.strip() for p in spec.split(",") if p.strip()]
    out = []
    for p in parts:
        if p.endswith("%"):
            try:
                v = float(p.rstrip("%"))
                out.append(page_width_cm * v / 100.0)
            except ValueError:
                out.append(0)
        else:
            try:
                out.append(float(p))
            except ValueError:
                out.append(0)
    # pad/truncate
    while len(out) < cols:
        out.append(0)
    return out[:cols]


# ---------- Top-level ----------

def generate_docx(article: Article, citation_style: str = "apa") -> bytes:
    doc = Document()
    _set_margins(doc)
    _register_styles(doc)

    j = article.journal
    _build_page_header(doc, j, article)
    _build_page_footer(doc)

    # Title — justified (rata kiri-kanan)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run(article.title or "Untitled")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Calibri"
    _suppress_hyphens(title_p)

    if article.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sr = sub.add_run(article.subtitle)
        sr.italic = True
        sr.font.size = Pt(12)
        sr.font.name = "Calibri"
        _suppress_hyphens(sub)

    # Authors
    if article.authors:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        names = []
        for a in article.authors:
            n = a.full_name or f"{a.given_name} {a.family_name}".strip()
            if a.corresponding:
                n += "*"
            names.append(n)
        names_run = ap.add_run(", ".join(names))
        names_run.font.size = Pt(11)
        names_run.font.name = "Times New Roman"
        names_run.bold = True

        for i, a in enumerate(article.authors, 1):
            if a.affiliation:
                af = doc.add_paragraph()
                af.alignment = WD_ALIGN_PARAGRAPH.LEFT
                afr = af.add_run(f"{i}. {a.affiliation}{(', ' + a.country) if a.country else ''}")
                afr.italic = True
                afr.font.size = Pt(9)
                afr.font.name = "Times New Roman"
                af.paragraph_format.space_after = Pt(0)

    figures_map = {f.id: f for f in (article.figures or [])}
    font_family = article.font_family or "Times New Roman"

    # Abstract — styled with shading + left border (matches PDF preview)
    if article.abstract.english:
        # ABSTRACT title
        ah = doc.add_paragraph(style="OpenJATS H2")
        ah_run = ah.add_run("ABSTRACT")
        ah_run.font.size = Pt(10)
        ah_run.bold = True
        _set_paragraph_shading(ah, "F1F5F9")
        _set_paragraph_left_border(ah, "475569", 24)
        # Body
        ap = doc.add_paragraph(style="OpenJATS Abstract")
        _add_inline_runs(ap, article.abstract.english, base_font=font_family, base_size=10)
        _set_paragraph_shading(ap, "F1F5F9")
        _set_paragraph_left_border(ap, "475569", 24)

    if article.keywords:
        kp = doc.add_paragraph(style="OpenJATS Abstract")
        kp.paragraph_format.first_line_indent = Cm(0)
        kr = kp.add_run("Keywords: ")
        kr.bold = True
        kr2 = kp.add_run(", ".join(article.keywords))
        kr2.italic = True
        _set_paragraph_shading(kp, "F1F5F9")
        _set_paragraph_left_border(kp, "475569", 24)

    history_bits = []
    if article.received_date:
        history_bits.append(f"Received: {article.received_date}")
    if article.revised_date:
        history_bits.append(f"Revised: {article.revised_date}")
    if article.accepted_date:
        history_bits.append(f"Accepted: {article.accepted_date}")
    if history_bits:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = hp.add_run(" · ".join(history_bits))
        hr.italic = True
        hr.font.size = Pt(9)
        hr.font.name = "Calibri"
        hp.paragraph_format.space_after = Pt(6)
        _set_paragraph_shading(hp, "F1F5F9")
        _set_paragraph_left_border(hp, "475569", 24)

    # Creative Commons License block — matches PDF "license-banner"
    lic_map = {
        "CC-BY 4.0": ("CC BY 4.0", "https://creativecommons.org/licenses/by/4.0/"),
        "CC-BY-SA 4.0": ("CC BY-SA 4.0", "https://creativecommons.org/licenses/by-sa/4.0/"),
        "CC-BY-NC 4.0": ("CC BY-NC 4.0", "https://creativecommons.org/licenses/by-nc/4.0/"),
        "CC-BY-NC-SA 4.0": ("CC BY-NC-SA 4.0", "https://creativecommons.org/licenses/by-nc-sa/4.0/"),
        "CC-BY-ND 4.0": ("CC BY-ND 4.0", "https://creativecommons.org/licenses/by-nd/4.0/"),
        "CC0": ("CC0 1.0 Universal", "https://creativecommons.org/publicdomain/zero/1.0/"),
        "All rights reserved": ("All rights reserved", ""),
    }
    if article.license and article.license in lic_map:
        lic_name, lic_url = lic_map[article.license]
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after = Pt(8)
        lp.paragraph_format.first_line_indent = Cm(0)
        r1 = lp.add_run(f"© {j.year or ''} The Author(s). ")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.name = "Calibri"
        r2 = lp.add_run(f"Published under {lic_name}")
        r2.font.size = Pt(9)
        r2.font.name = "Calibri"
        if lic_url:
            r3 = lp.add_run(f" ({lic_url})")
            r3.font.size = Pt(9)
            r3.font.name = "Calibri"
            r3.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)
        r4 = lp.add_run(". This is an open access article distributed under the terms of the license, which permits unrestricted use, distribution and reproduction provided the original work is properly cited.")
        r4.font.size = Pt(9)
        r4.font.name = "Calibri"
        r4.italic = True
        _set_paragraph_shading(lp, "F1F5F9")
        _set_paragraph_left_border(lp, "0EA5E9", 24)

    if article.abstract.indonesian:
        _heading(doc, "Abstrak", level=1)
        _add_body_para(doc, article.abstract.indonesian, style="OpenJATS Abstract", indent=False, font_family=font_family)

    sec_map = [
        ("Introduction", article.sections.introduction),
        ("Methods", article.sections.methods),
        ("Results", article.sections.results),
        ("Discussion", article.sections.discussion),
        ("Conclusion", article.sections.conclusion),
    ]
    for title, body in sec_map:
        if body and body.strip():
            _heading(doc, title, level=1)
            _render_body(doc, body, figures_map, font_family=font_family)

    back_map = [
        ("Acknowledgements", article.sections.acknowledgement),
        ("Funding", article.sections.funding),
        ("Conflict of Interest", article.sections.conflict_of_interest),
        ("Data Availability", article.sections.data_availability),
        ("Author Contributions", article.sections.author_contributions),
        ("Ethical Approval", article.sections.ethical_approval),
    ]
    for title, body in back_map:
        if body and body.strip():
            _heading(doc, title, level=2)
            _add_body_para(doc, body, style="OpenJATS Body", indent=False, font_family=font_family)

    if article.references:
        _heading(doc, "References", level=1)
        formatted = format_references(article.references, citation_style)
        for line in formatted:
            p = doc.add_paragraph(style="OpenJATS Reference")
            # Parse "*text*" segments as italic runs
            segments = re.split(r"\*([^*]+)\*", line)
            for idx, seg in enumerate(segments):
                if seg == "":
                    continue
                run = p.add_run(seg)
                if idx % 2 == 1:  # odd index = inside *…*
                    run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
